# -*- coding: utf-8 -*-
"""Renderiza el dossier a DOCX (editable) y PDF."""
import pathlib
import sys

sys.path.insert(0, str(pathlib.Path(__file__).parent))
from contenido import CONTENIDO  # noqa: E402

SALIDA = pathlib.Path("C:/DevsJesus/tuvetia/skeleton/docs/entrega")
SALIDA.mkdir(parents=True, exist_ok=True)

C = list(CONTENIDO)
A = C.append

# ---------------------------------------------------------------------------------------------
# Corte final del 16 de agosto: encabezado, semanas 5-7 y adendas.
# contenido.py quedó congelado al 1 de agosto; los ajustes se aplican aquí, sobre la copia.
# ---------------------------------------------------------------------------------------------
for _i, (_t, _v) in enumerate(C):
    if _t == "p" and _v.startswith("Contrato COT-2026-TUV-001"):
        C[_i] = ("p", "Contrato COT-2026-TUV-001 · TUVET IA · Corte: 31 de julio de 2026 · "
                      "Adenda: 3 de agosto · CORTE FINAL: 16 de agosto de 2026")
        break

SEMANAS_NUEVAS = [
    ("h3", "Semana 5 — WhatsApp con tráfico real, calendario multi-proveedor "
           "(31 de julio al 3 de agosto)"),
    ("code", "2d1a7e5   31-jul   Bandeja en TIEMPO REAL: los mensajes entran al instante, no cada 15 s"),
    ("code", "57005b1   01-ago   La cascada registraba el modelo equivocado — la traza ya dice la verdad"),
    ("code", "1450614   02-ago   El modo auto de WhatsApp puede mirar la agenda y proponer citas"),
    ("code", "e9f9a80   02-ago   Permiso de ADMIN exigido en las tres acciones que salen de la clínica sin vuelta atrás"),
    ("code", "26dbb02   03-ago   Google Calendar por Composio, en vez de OAuth propio"),
    ("code", "767152f   03-ago   Outlook Calendar por Composio — una sola ruta para los dos proveedores"),
    ("code", "4e0f5ca   03-ago   El webhook de WhatsApp ya no descarta mensajes en silencio"),
    ("code", "53314eb   03-ago   Fotos en el hilo de WhatsApp: se ven, en orden, y /f/<token> ya no da 404"),
    ("code", "fdee418   03-ago   Un envío que falla lo dice en 20 segundos, no en 40"),
    ("code", "470a5a2   03-ago   Mandarse un mensaje al número de la propia clínica daba un 400 sin explicación"),
    ("code", "fc54fb8   03-ago   Normalización E.164: el número de la ficha sale con indicativo, y el error del proveedor se traduce"),
    ("p", "El 3 de agosto WhatsApp queda operando con tráfico real bidireccional en producción, "
          "tras corregir la configuración del servicio de transporte. El detalle, con su evidencia "
          "y sus límites, está en la Adenda del final."),

    ("h3", "Semana 6 — correo unificado, el asistente en toda la app y el rediseño "
           "(3 al 12 de agosto)"),
    ("code", "7a2f115   03-ago   El health no miraba el correo, la tarjeta mentía y un envío fallido no dejaba rastro"),
    ("code", "30490d6   03-ago   La cobranza vuelve a leer las respuestas, ahora por Composio — IMAP se retira"),
    ("code", "5a6a370   03-ago   Athos alcanzable desde cualquier pantalla, sabiendo qué estás mirando"),
    ("code", "aaf9d28   03-ago   Cada propuesta registra de dónde vino, y cuánto cuesta cada superficie de IA"),
    ("code", "80046f4   11-ago   La invitación de equipo se envía por Resend, con botón propio"),
    ("code", "4fa7f64   11-ago   La grabación de una consulta ya no muere al cambiar de pantalla"),
    ("code", "a8de57f   11-ago   Rediseño: blanco y menta, y la barra lateral se parte en consultorio y CRM"),
    ("code", "5909a43   11-ago   Onboarding: un riel que dice qué le falta a la clínica"),
    ("code", "bd5ffec   12-ago   Athos primero — la app abre en la conversación, con la clínica al lado"),
    ("code", "e94c114   12-ago   El chat del asistente toma la forma familiar de ChatGPT"),

    ("h3", "Semana 7 — la consulta en vivo, seguridad y topes (15 y 16 de agosto)"),
    ("code", "462ed9e   15-ago   El cuaderno: durante la consulta el vet no tenía dónde escribir NADA"),
    ("code", "aae70ef   15-ago   «Iniciar consulta» ahora inicia la consulta: grabación, transcripto y cuaderno juntos"),
    ("code", "50e5216   15-ago   La alerta de alergia nombra el fármaco, y lo dice en el plan de la nota"),
    ("code", "8656e7a   16-ago   La nota decía «evidencia suficiente» cuando el juez dijo lo contrario"),
    ("code", "efeeb95   15-ago   La app no tenía ninguna red cuando algo se rompe: páginas de error en tres niveles"),
    ("code", "21d6eb5   15-ago   El modelo podía falsificar la marca que existe para controlarlo"),
    ("code", "dac3f1d   15-ago   Tope de gasto de IA por clínica, persistente entre lambdas"),
    ("code", "e67bd2e   15-ago   El vet ve cuánto cupo de IA le queda antes de topárselo"),
    ("code", "dea56b6   15-ago   Desactivar una cuenta la desactiva: gate en la base (migraciones 0059-0061)"),
    ("code", "5b6ff86   16-ago   El interruptor de desactivación en el panel de administración"),
    ("code", "5efed7f   16-ago   Desactivar también corta las APIs, no sólo las pantallas"),
    ("code", "ab80de5   16-ago   Un paciente creado ya se puede corregir: edición de la ficha"),
    ("code", "17a4caa   16-ago   El pago que el cliente ya entregó no puede perderse"),
    ("code", "de6b825   16-ago   El asistente sabe qué pantalla tenés delante — y quién espera en la clínica (d56a216)"),
    ("p", "El detalle de esta semana, con su evidencia y sus límites, está en la Segunda adenda "
          "del final."),
]
_idx = next(i for i, (t, v) in enumerate(C) if t == "h1" and v.startswith("Evidencia"))
C[_idx:_idx] = SEMANAS_NUEVAS

# ---------------------------------------------------------------------------------------------
# Adenda — corte del 3 de agosto
# ---------------------------------------------------------------------------------------------
A(("h1", "Adenda — corte del 3 de agosto de 2026"))
A(("p", "Lo de abajo es evidencia posterior al corte del 31 de julio, tomada el 3 de agosto sobre "
        "producción con los mismos criterios del resto del documento: se dice qué clase de "
        "evidencia respalda cada afirmación, y dónde está el límite."))

A(("h2", "Adenda A — WhatsApp: transporte verificado con tráfico real"))
A(("p", "QUÉ ESTABA FALLANDO. El transporte de WhatsApp (Evolution, elegido el 28 de julio con "
        "consentimiento registrado de la clínica) corría sin su cache configurado: la sesión "
        "figuraba «conectada» pero no procesaba ni un mensaje — cero filas en la tabla de mensajes "
        "en toda su historia, y cero también en la base interna del propio transporte. El "
        "diagnóstico se hizo por capas y con evidencia en cada una: webhook registrado con URL y "
        "token correctos (probado con una petición real → 200), ruta desplegada, y el hueco "
        "aislado aguas arriba, en el servicio de transporte."))
A(("p", "LA CORRECCIÓN. Configuración del servicio (3 de agosto), no código. El mismo día quedó "
        "fluyendo tráfico real bidireccional: mensajes entrantes de un teléfono real, respuestas "
        "desde la bandeja, y la sincronización de lo que el veterinario escribe desde su propio "
        "teléfono (evidencia de estado: filas inbound y outbound del 3 de agosto en "
        "whatsapp_messages, con cero fallidas de transporte)."))
A(("p", "QUÉ MÁS QUEDÓ CERRADO EL MISMO DÍA (evidencia documental): el webhook ya no puede "
        "descartar mensajes en silencio — cada evento deja una línea de log con su motivo "
        "(4e0f5ca); las fotos se ven en el hilo, en orden, y la ruta /f/<token> existe (53314eb); "
        "un envío que falla lo dice en ~20 segundos con un mensaje que explica el motivo "
        "(fdee418, d632fb3)."))
A(("h3", "Límites de la evidencia"))
A(("quote", "Los envíos que parten de la ficha del titular fallaban si el teléfono estaba guardado "
            "sin indicativo de país: el 3 de agosto quedaron dos acciones failed con el error "
            "íntegro y verificable en athos_actions.error (el transporte responde exists:false "
            "para el número sin 57). QUEDÓ CORREGIDO ESE MISMO DÍA: la normalización E.164 en el "
            "único camino de salida, y el error crudo del proveedor traducido a un mensaje que "
            "explica el motivo (fc54fb8). Las respuestas desde la bandeja nunca lo necesitaron "
            "porque responden a números que llegaron por webhook, ya completos."))
A(("quote", "El botón «Sugerir» de la bandeja registró su primer uso real el 3 de agosto sin "
            "producir propuesta; está en diagnóstico (la cascada del agente resolvió a un modelo "
            "distinto del previsto para superficies con herramientas)."))
A(("quote", "El modo auto responde únicamente saludos, horarios y citas; ante cualquier cosa "
            "clínica CALLA POR DISEÑO y el mensaje queda para el veterinario. Además arranca con "
            "rampa anti-baneo (5 respuestas/día el primer día). Un silencio ante una pregunta "
            "clínica es la salvaguarda operando, no un defecto."))

A(("h2", "Adenda B — Calendario: de OAuth propio a Composio, y ahora también Outlook"))
A(("p", "El 3 de agosto la integración de calendario migró a Composio (26dbb02) y se añadió "
        "Outlook Calendar por la misma ruta (767152f): un solo camino de código para los dos "
        "proveedores. Quedaron corregidos además tres defectos de ruteo: la cita va al calendario "
        "del ADMINISTRADOR con el veterinario invitado (3e96aa8), la pantalla de conexiones "
        "muestra el calendario DE LA CLÍNICA y no el de quien mira (23b116f), y una cita que no "
        "llega al calendario AVISA en vez de fallar en silencio (032982d). La conexión es por "
        "miembro, no por clínica: cada uno usa su propia cuenta."))

A(("h2", "Adenda C — Verificación pre-demo del 3 de agosto (evidencia de estado y comportamiento)"))
A(("p", "Verificación completa de producción la mañana del 3 de agosto:"))
A(("tabla", [["Qué", "Resultado"],
             ["Backend /health", "200 en 0,5 s"],
             ["Endpoints protegidos sin credencial", "401 (el gate de auth opera)"],
             ["CORS desde el dominio del frontend", "aceptado para el origen exacto"],
             ["Pipeline RAG completo (petición real autenticada)",
              "200 con evidence_level y 8 fragmentos con fuente y localizador"],
             ["Latencia del pipeline RAG", "6,5 s en caliente · 19,9 s la primera consulta en frío"],
             ["Corpus en producción", "~519.900 fragmentos indexados"],
             ["Datos de demostración en la cuenta del cliente",
              "22 pacientes, 24 notas (7 borradores por aprobar), 11 consultas listas para el "
              "Modo Fantasma, 8 alergias (3 severas), memoria de paciente indexada"]]))
A(("quote", "Límite: la cifra de latencia en frío importa operativamente — la primera consulta "
            "tras un rato de inactividad tarda ~20 s. Para una demostración conviene hacer una "
            "consulta de calentamiento minutos antes."))

# ---------------------------------------------------------------------------------------------
# Segunda adenda — corte final del 16 de agosto
# ---------------------------------------------------------------------------------------------
A(("h1", "Segunda adenda — corte final del 16 de agosto de 2026"))
A(("p", "Trabajo posterior al 3 de agosto, con la misma regla del resto del documento: cada "
        "afirmación dice qué clase de evidencia la respalda, y dónde está el límite."))

A(("h2", "Adenda D — La consulta en vivo: una sola superficie para grabar, escribir y firmar"))
A(("p", "Hasta el 15 de agosto la consulta estaba repartida: TRES superficies de grabación "
        "distintas, DOS cuadernos que no se hablaban entre sí, y ningún lugar donde el veterinario "
        "escribiera durante la consulta. El 15 de agosto se unificó (bb8a1d4, 84ee87c, aae70ef): "
        "«iniciar consulta» inicia la grabación con el transcripto en vivo y el cuaderno lado a "
        "lado, la grabación sobrevive al cambio de pantalla (4fa7f64), y el cuaderno persiste en "
        "la base (migración 0058, con pruebas propias en el servicio y en el front). De 3-4 clics "
        "entre «quiero grabar» y «grabando» quedaron 2."))
A(("p", "Además la nota ganó dos garantías (evidencia documental): la alerta de alergia NOMBRA EL "
        "FÁRMACO y lo dice en el plan, que es donde el vet decide (50e5216); y la nota ya no puede "
        "decir «evidencia suficiente» cuando el juez de abstención dictaminó lo contrario "
        "(8656e7a) — la etiqueta sale del veredicto real, no de un texto fijo."))

A(("h2", "Adenda E — Seguridad y antifraude: la fase 1 operando, y lo que falta dicho de frente"))
A(("p", "Lo construido entre el 15 y el 16 de agosto (evidencia documental y de estado):"))
A(("b", "Desactivar una cuenta la desactiva. El gate vive en la base (migraciones 0059-0061), "
        "verificado contra el Postgres de producción con 7 comprobaciones en OK. El bypass por "
        "PATCH del propio usuario quedó cerrado (93b6aac, 0060), las APIs también se cortan —no "
        "sólo las pantallas— (5efed7f), y el panel de administración tiene el interruptor "
        "(5b6ff86)."))
A(("b", "El modelo no puede falsificar la marca que existe para controlarlo (21d6eb5): la tarjeta "
        "de aprobación es del código, no del texto del modelo."))
A(("b", "Permiso de ADMIN exigido en las tres acciones que salen de la clínica sin vuelta atrás "
        "(e9f9a80, del 2 de agosto)."))
A(("b", "El panel de administración muestra el costo real por tokens (/admin/costos) y los "
        "usuarios con su última fecha de ingreso, incluida la señal «nunca entró»."))
A(("h3", "Límites de la evidencia"))
A(("quote", "La desactivación automática NO EXISTE: no hay ningún cron que desactive por señales. "
            "Hoy es un interruptor manual."))
A(("quote", "Las señales de abuso (correos parecidos, direcciones IP) no están construidas. El "
            "plan completo está en docs/ANTIFRAUDE-2026-08-15.md, que documenta también qué "
            "señales dependen de la retención de logs de autenticación."))

A(("h2", "Adenda F — El tope de gasto de IA: construido, y hoy apagado a propósito"))
A(("p", "El 15 de agosto se construyó el tope mensual de gasto de IA por clínica (dac3f1d, "
        "6f3f8b6): persiste en la base —sobrevive a la siguiente lambda—, las SEIS superficies que "
        "consumen IA pasan por él (chat, bandeja, modo automático, cartera y las dos de visión), y "
        "el veterinario VE CUÁNTO LE QUEDA antes de topárselo (e67bd2e), en vez de un corte a "
        "ciegas."))
A(("h3", "Límite de la evidencia"))
A(("quote", "Verificado el 16 de agosto contra las variables de producción de Vercel: "
            "ATHOS_TOPE_MENSUAL_POR_CLINICA no está configurada, así que hoy el tope NO CORTA "
            "NADA. Es deliberado — el acta tiene abierta la decisión del límite exacto entre el "
            "plan gratuito y el de pago — pero mientras esa decisión no se tome, el único freno "
            "vivo es el límite de peticiones por minuto. Encenderlo es poner una variable, sin "
            "redespliegue de código."))

A(("h2", "Adenda G — El rediseño, el asistente en el centro, y la salud del código al corte final"))
A(("p", "Entre el 11 y el 16 de agosto la interfaz cambió de forma (evidencia documental): blanco "
        "y menta con la barra partida en consultorio y CRM (a8de57f), la app abre en la "
        "conversación con Athos y la clínica al lado (bd5ffec), el chat toma la forma familiar de "
        "ChatGPT (e94c114), tab bar en el teléfono (87d2fed), páginas de error en tres niveles "
        "donde antes no había ninguna (efeeb95), y un riel de onboarding que dice qué le falta a "
        "la clínica (5909a43). El asistente además SABE QUÉ PANTALLA TENÉS DELANTE (de6b825) y "
        "QUIÉN ESTÁ ESPERANDO EN LA CLÍNICA, sin gastar un token (d56a216)."))
A(("p", "Salud del código al corte, medida el 16 de agosto (docs/DIAGNOSTICO-2026-08-16.md):"))
A(("tabla", [["Qué", "Resultado"],
             ["Errores de tipos (tsc --noEmit)", "0"],
             ["Avisos de lint (eslint --max-warnings=0)", "0"],
             ["Pruebas del front", "743 en 69 archivos, todas en verde"],
             ["Pruebas del servicio", "261 en verde"],
             ["Build de producción", "OK"],
             ["Rutas del dashboard sin enlace entrante", "0 de 32"]]))
A(("h3", "Límites de la evidencia"))
A(("quote", "La medición es del commit aae70ef (15-ago); los commits del 16 añadieron MÁS pruebas "
            "(entre ellas, la primera suite de la ruta que ejecuta acciones: 32d72c2), así que el "
            "número real al corte es mayor, no menor."))
A(("quote", "Lo VISUAL del rediseño (radios, espaciados, el layout de la consulta) no se verificó "
            "a ojo en esta máquina: está cubierto por tipos, lint, pruebas y build, y debe mirarse "
            "en el despliegue de producción."))
A(("quote", "El botón «Sugerir» de la bandeja de WhatsApp sigue EN DIAGNÓSTICO desde el 3 de "
            "agosto (ver Adenda A). Desde aaf9d28 cada propuesta registra su procedencia y su "
            "costo, que es el instrumento para ese diagnóstico."))
A(("quote", "El CORREO migró: la lectura de respuestas de cobranza ya no es IMAP sino Composio "
            "(30490d6, 3-ago). El Punto 6 describe el estado al 31 de julio; la arquitectura "
            "vigente es la de Composio."))

# ---------------------------------------------------------------------------------------------
# El método: por qué el proyecto avanzó en ese orden
# ---------------------------------------------------------------------------------------------
A(("h1", "El método — por qué se avanzó en ese orden"))
A(("p", "Las fechas de la sección anterior no son casualidad. Responden a cuatro decisiones "
        "técnicas que conviene explicar, porque son las que justifican por qué una función podía "
        "estar en producción el 16 de julio y aun así necesitar dos semanas más de trabajo."))

A(("h2", "1. Primero lo determinístico, después la inteligencia artificial"))
A(("p", "La arquitectura parte de una regla: gastar la menor inteligencia artificial posible. El "
        "buscador que recupera literatura es determinístico y no consume tokens; la IA sólo "
        "interviene en dos puntos — entender la consulta y redactar la respuesta citada."))
A(("p", "Por eso la primera versión de la abstención (13 de julio) era un umbral numérico sobre el "
        "score del buscador: era lo barato, lo reproducible y lo que no dependía de ningún modelo. "
        "Sólo cuando se midió y se comprobó que no discriminaba se pagó el costo de un juez que "
        "LEE los pasajes."))
A(("quote", "Ese orden es deliberado, no un descuido. Introducir un modelo de lenguaje donde una "
            "regla determinística alcanza añade costo, latencia y una fuente de error que no se "
            "puede auditar. La secuencia correcta es intentar lo barato, medirlo, y escalar sólo "
            "si los datos lo exigen."))

A(("h2", "2. Las reglas duras viven en el código, no en el prompt"))
A(("p", "Está medido en este proyecto: pedirle al modelo por prompt que no dé una dosis sin peso "
        "del paciente falla en 2 de 23 casos, y con un prompt más resolutivo empeora a 9 de 23 "
        "—porque pedirle que decida lo empuja a decidir—. Por eso las garantías críticas están "
        "impuestas por código:"))
A(("b", "El gate de alergia severa se calcula desde la tabla de alergias, no lo decide el modelo."))
A(("b", "El tapado de dosis cuando faltan datos lo hace un guard, no el prompt."))
A(("b", "La verificación de citas descarta referencias que no correspondan a un documento "
        "recuperado, sin preguntarle al modelo."))
A(("b", "El aviso de evidencia limitada lo emite el código cuando el juez dictamina esa banda."))
A(("p", "Esto explica por qué varias correcciones de las últimas semanas fueron mover una regla "
        "desde el prompt hacia el código: no es refactorización cosmética, es la diferencia entre "
        "una garantía y una sugerencia."))

A(("h2", "3. Medir antes de construir — y arreglar el instrumento cuando está mal"))
A(("p", "Buena parte del trabajo del 28 al 31 de julio no fue escribir funciones nuevas sino "
        "construir con qué medir las existentes. Ese esfuerzo devolvió tres hallazgos que ninguna "
        "revisión de código habría encontrado:"))
A(("b", "El umbral de abstención daba «hay evidencia» en 187 de 187 casos: la función existía pero "
        "su criterio estaba saturado."))
A(("b", "18 de 24 respuestas citaban al menos un pasaje que no respaldaba la afirmación: la "
        "procedencia era correcta, la pertinencia no."))
A(("b", "24 de los 42 casos negativos del banco de pruebas NO eran negativos: el instrumento de "
        "medición estaba roto, y toda cifra anterior medía otra cosa."))
A(("quote", "El tercero es el más incómodo y el más importante. Descubrir que el propio "
            "instrumento estaba mal obliga a descartar mediciones ya reportadas. Se hizo, y por "
            "eso las cifras de este documento son las que se sostienen."))

A(("h2", "4. Lo que se puede contar, se cuenta"))
A(("p", "Se probó usar un modelo de lenguaje como juez de calidad y se midió su fiabilidad: tiene "
        "un ruido de ±7 sobre 40 puntos entre corridas idénticas, y un sesgo de posición del 78% "
        "hacia la primera opción mostrada. Es decir, el juez podía «demostrar» una mejora que no "
        "existía."))
A(("p", "Desde entonces la regla es: si una propiedad se puede contar de forma determinística, se "
        "cuenta. La abstención se mide contra un hecho comprobable —si algún pasaje recuperado "
        "está indexado con el descriptor de la consulta en el árbol MeSH—, la transcripción contra "
        "una verdad de terreno conocida, y la proporcionalidad de las respuestas contando señales "
        "clínicas y caracteres. Ninguna cifra de este documento sale de la opinión de un modelo."))

# ---------------------------------------------------------------------------------------------
# Guía para grabar los demos probatorios
# ---------------------------------------------------------------------------------------------
A(("h1", "Guía para grabar los demos probatorios"))
A(("p", "Un video vale como evidencia sólo si es imposible discutir lo que muestra. Estas reglas "
        "existen para que nadie pueda decir «eso está editado» o «así no se mide»."))

A(("h2", "Reglas que valen para todos los videos"))
A(("b", "Una sola toma, sin cortes. Si algo sale mal, se repite el video entero. Un corte invalida "
        "la evidencia."))
A(("b", "Reloj visible en pantalla desde el primer segundo. Abrí time.is en una ventana lateral: "
        "muestra la hora oficial con segundos y sirve de sello temporal verificable."))
A(("b", "Mostrá la URL. Que se vea que es tuvetia.vercel.app y no un entorno local."))
A(("b", "Narrá lo que vas a hacer ANTES de hacerlo. «Voy a preguntar X y espero que pase Y.» Una "
        "predicción cumplida es más creíble que una explicación posterior."))
A(("b", "No edites nada. Ni recortes, ni acelerados, ni música. El archivo crudo."))
A(("b", "Nombrá el archivo con punto y fecha: `P05-latencia-2026-08-01.mp4`."))

A(("h2", "Con qué grabar"))
A(("p", "En Windows, la Xbox Game Bar viene incluida: tecla Windows + G, y grabás con Windows + "
        "Alt + R. Si necesitás mejor calidad o mostrar el cronómetro con precisión, OBS Studio es "
        "gratuito y permite superponer un cronómetro sobre la pantalla."))
A(("p", "Para el punto de latencia conviene además grabar con el panel de red del navegador "
        "abierto (F12, pestaña Red): ahí queda el tiempo real de cada petición, que es un dato "
        "que el propio navegador certifica."))

A(("h2", "Video por video: qué mostrar exactamente"))

VIDEOS = [
    ("P05 · Latencia", "2 a 3 minutos", [
        "Mostrá time.is y la URL de producción.",
        "Abrí F12 en la pestaña Red y limpiá el registro.",
        "Escribí una consulta clínica real y decí en voz alta: «voy a enviar ahora».",
        "Enviá y NO toques nada. Dejá que se vea aparecer el primer texto.",
        "Al aparecer el primer token, decilo en voz alta: «primer token».",
        "Cuando termine, mostrá en el panel de red el tiempo total de la petición.",
        "Cerrá mostrando el reloj otra vez.",
    ], "Lo que prueba: que el primer token llega en torno a 12,8 s y la respuesta completa en "
       "27,6 s. Los cinco minutos reportados no se reproducen."),

    ("P03 · Abstención", "3 a 4 minutos", [
        "Primero una consulta que SÍ tenga literatura (por ejemplo, dermatitis atópica canina). "
        "Mostrá que responde citando fuentes.",
        "Después una consulta sin cobertura real en el corpus. Anunciá antes: «esto no debería "
        "tener evidencia suficiente».",
        "Mostrá el aviso de evidencia limitada o la abstención en pantalla.",
        "Si aparece la advertencia de evidencia limitada, leela en voz alta.",
    ], "Lo que prueba: que el sistema distingue y lo declara. Es el contraste entre los dos casos "
       "lo que convence, no el segundo caso solo."),

    ("P04 · Citas correctas", "2 minutos", [
        "Hacé una consulta que devuelva fuentes.",
        "Hacé clic en una fuente y mostrá que abre el artículo real (PubMed o el enlace del "
        "documento).",
        "Volvé y abrí otra. Mostrá que el título y el año coinciden con lo que abre.",
    ], "Lo que prueba: que las fuentes existen y son las recuperadas. El argumento fuerte es "
       "estructural —el modelo sólo emite un número, el resto lo reconstruye el código— pero verlo "
       "abrir el artículo real lo hace tangible."),

    ("P08 · Transcripción en vivo", "3 a 5 minutos", [
        "IMPORTANTE: grabá con DOS personas hablando, no una sola. Es lo único que demuestra la "
        "separación de hablantes.",
        "Mostrá la pantalla de consentimiento antes de grabar (Ley 1581).",
        "Iniciá la grabación y hablá alternando: uno hace de dueño («Doctor, mi perro no come») y "
        "otro de veterinario («vamos a palpar el abdomen»).",
        "Mostrá el texto apareciendo EN VIVO mientras hablan.",
        "Al detener, mostrá que los roles quedaron bien asignados: Titular y Veterinario.",
        "Mostrá la fecha de la consulta y comparala con el reloj en pantalla.",
    ], "Lo que prueba: los tres defectos cerrados de una sola vez — en vivo, roles correctos y "
       "fecha correcta. Este es el video más valioso de todos, y el que cubre el hueco que la "
       "medición sintética no pudo cubrir."),

    ("P01 · Cascada de tres modelos", "3 minutos", [
        "Este NO se puede demostrar desde la interfaz: el fallback sólo se activa si el proveedor "
        "principal falla.",
        "Grabá en su lugar la ejecución de la prueba de caída forzada en la terminal, mostrando "
        "que responde DeepSeek, luego Gemini y luego Claude.",
        "Mostrá también las variables de entorno de Railway con las tres claves configuradas "
        "(sin revelar los valores: basta con que se vean los nombres).",
    ], "Lo que prueba: que la cascada existe y funciona. Sé explícito en el video sobre que es una "
       "prueba de caída forzada y no tráfico real: es lo que hace creíble todo lo demás."),

    ("P07 · Google Calendar", "2 a 3 minutos", [
        "Mostrá el calendario de la plataforma con las citas.",
        "Abrí Google Calendar en otra pestaña y mostrá las mismas citas.",
        "Creá una cita en la plataforma y mostrá cómo aparece en Google.",
    ], "Lo que prueba: la sincronización operando. Hay 11.550 citas sincronizadas, así que el "
       "volumen habla solo."),

    ("P09 · Invitaciones", "2 minutos", [
        "Usá un correo que NO tenga cuenta en la plataforma — es el caso que estaba roto.",
        "Enviá la invitación desde ajustes de equipo.",
        "Abrí el correo, hacé clic en el enlace.",
        "Mostrá que aterriza en la invitación y que al aceptar entra a la clínica.",
    ], "Lo que prueba: el quinto defecto corregido. Con un correo que ya tenga cuenta NO se "
       "demuestra nada: ése es el camino que siempre funcionó."),

    ("P10 · Historial y memoria", "2 minutos", [
        "Escribí una pregunta al asistente y esperá la respuesta.",
        "RECARGÁ la página por completo (F5).",
        "Mostrá que la conversación anterior sigue ahí.",
        "Preguntá: «¿cuál fue la última pregunta que te hice?» y mostrá que responde bien.",
    ], "Lo que prueba: que el historial persiste entre sesiones. La recarga es la parte "
       "importante: sin ella no se demuestra persistencia."),

    ("P06b · WhatsApp en vivo", "2 a 3 minutos", [
        "Mostrá la bandeja de Comunicaciones en la URL de producción, con el reloj visible.",
        "Desde un teléfono, enviá un WhatsApp al número conectado de la clínica y decí antes qué "
        "vas a escribir.",
        "Mostrá el mensaje apareciendo en la bandeja AL INSTANTE (sin recargar: la bandeja es en "
        "tiempo real).",
        "Respondé desde la bandeja y mostrá el teléfono recibiendo la respuesta.",
        "Mandá una FOTO desde el teléfono y mostrá que se ve en el hilo.",
        "Escribí un mensaje desde el teléfono de la clínica (no desde la plataforma) y mostrá que "
        "también aparece en el hilo: es la sincronización completa del número.",
    ], "Lo que prueba: el ciclo entero de WhatsApp operando en producción — entrante en tiempo "
       "real, respuesta, media y sincronización del teléfono propio. Es la evidencia de "
       "comportamiento que respalda la Adenda A."),
]

for titulo, duracion, pasos, prueba in VIDEOS:
    A(("h3", f"{titulo}  ({duracion})"))
    for p in pasos:
        A(("b", p))
    A(("quote", prueba))
    A(("hueco", "Espacio para pegar el video o su captura"))

A(("h2", "Qué NO hacer"))
A(("b", "No grabes en un entorno local. Si la URL no es la de producción, no prueba nada."))
A(("b", "No uses datos de pacientes reales de una clínica cliente sin autorización. Creá pacientes "
        "de demostración."))
A(("b", "No repitas la toma hasta que «salga bien» y muestres sólo esa. Si el sistema falla en un "
        "intento, eso también es información — y si se descubre después, cuesta mucho más caro."))
A(("b", "No afirmes en el video nada que el video no muestre. Si querés dar contexto, decilo como "
        "contexto, no como demostración."))

# --- anexo de comandos -----------------------------------------------------------------------
A(("h1", "Anexo — comandos de verificación"))
A(("p", "Cualquiera con acceso al repositorio puede reproducir esta evidencia."))

A(("h3", "Verificar un commit concreto"))
A(("code", "git show cca7b87        # primer commit: citas + abstención + historial\n"
           "git log --format='%h %ad %s' --date=short --reverse | head -20"))

A(("h3", "Correr la suite completa"))
A(("code", "python scripts/auditoria.py        # backend + front + tipos + lint + build"))

A(("h3", "Medir la abstención contra producción"))
A(("code", "cd athos-service\n"
           "python scripts/calidad/abstencion_verdad.py --n 0"))

A(("h3", "Medir la transcripción en vivo contra Deepgram"))
A(("code", "python scripts/calidad/transcripcion_vivo_verificar.py audio.wav referencia.txt"))

A(("h3", "Correr el banco de calidad del agente"))
A(("code", "RUN_BANCO=1 ANTHROPIC_API_KEY=... \\\n"
           "  npx vitest run --config vitest.e2e.config.mts e2e/banco-agente.e2e.ts"))

A(("h3", "Comprobar la configuración de producción"))
A(("code", 'curl -H "Authorization: Bearer $CRON_SECRET" \\\n'
           "  https://tuvetia.vercel.app/api/health"))

A(("h2", "Documentos de referencia en el repositorio"))
for d in [
    "docs/SOPORTES-MILESTONE2-2026-07-30.md — respuesta punto por punto a la guía de soportes",
    "docs/VERIFICACION-10-PUNTOS-2026-07-30.md — estado de los 10 puntos priorizados",
    "docs/ABSTENCION-MEDICION-2026-07-30.md — cómo se mide la abstención y por qué",
    "docs/BANCO-AGENTE-RESULTADO.md — banco de calidad del agente con las respuestas íntegras",
    "docs/COMPARATIVA-MODELOS-2026-07-30.md — prueba de caída forzada entre los tres proveedores",
    "docs/AGENT-SMOKE-TESTING.md — resultados del smoke testing",
    "docs/CONFIGURACION-PRODUCCION.md — qué variable vive dónde",
    "INVENTARIO-COMPONENTES.md — inventario formal de componentes",
    "docs/entrega/CAPA-AGENTICA-ESTADO.md — las 21 habilidades del asistente, cada una con su "
    "frase de prueba y el resultado esperado",
    "docs/EVOLUTION.md — operación del transporte de WhatsApp (decisión, riesgos y consentimiento)",
    "WHATSAPP.md — arquitectura de la capa de WhatsApp multi-proveedor",
    "FUNCIONALIDADES.md — mapa completo de funcionalidades, con el costo de operación de cada una",
    "docs/ANTIFRAUDE-2026-08-15.md — el plan antifraude: qué está construido y qué señales faltan",
    "docs/DIAGNOSTICO-2026-08-16.md — diagnóstico de cableado, UI y base de datos al corte final",
]:
    A(("b", d))

# ---------------------------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------------------------
from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Pt, RGBColor  # noqa: E402

doc = Document()
for st, sz in (("Normal", 10.5),):
    doc.styles[st].font.size = Pt(sz)
    doc.styles[st].font.name = "Calibri"

for tipo, val in C:
    if tipo == "h1":
        doc.add_page_break() if len(doc.paragraphs) > 3 else None
        doc.add_heading(val, level=0)
    elif tipo == "h2":
        doc.add_heading(val, level=1)
    elif tipo == "h3":
        doc.add_heading(val, level=2)
    elif tipo == "p":
        doc.add_paragraph(val)
    elif tipo == "b":
        doc.add_paragraph(val, style="List Bullet")
    elif tipo == "quote":
        p = doc.add_paragraph(val)
        p.paragraph_format.left_indent = Pt(24)
        for r in p.runs:
            r.italic = True
            r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    elif tipo == "code":
        p = doc.add_paragraph()
        r = p.add_run(val)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        p.paragraph_format.left_indent = Pt(18)
    elif tipo == "tabla":
        t = doc.add_table(rows=0, cols=len(val[0]))
        t.style = "Light Grid Accent 1"
        for i, fila in enumerate(val):
            celdas = t.add_row().cells
            for j, txt in enumerate(fila):
                celdas[j].text = str(txt)
                if i == 0:
                    for pr in celdas[j].paragraphs:
                        for run in pr.runs:
                            run.bold = True
        doc.add_paragraph()
    elif tipo == "hueco":
        p = doc.add_paragraph(f"[ {val} ]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.italic = True
            r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        doc.add_paragraph()

ruta_docx = SALIDA / "DOSSIER-EVIDENCIAS-MILESTONE2.docx"
doc.save(ruta_docx)
print(f"DOCX -> {ruta_docx}")

# ---------------------------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------------------------
from reportlab.lib import colors  # noqa: E402
from reportlab.lib.enums import TA_CENTER  # noqa: E402
from reportlab.lib.pagesizes import LETTER  # noqa: E402
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # noqa: E402
from reportlab.lib.units import cm  # noqa: E402
from reportlab.platypus import (PageBreak, Paragraph, SimpleDocTemplate,  # noqa: E402
                                Spacer, Table, TableStyle)

ss = getSampleStyleSheet()
E = {
    "h1": ParagraphStyle("H1", parent=ss["Title"], fontSize=19, spaceAfter=14, spaceBefore=6,
                         textColor=colors.HexColor("#12313f")),
    "h2": ParagraphStyle("H2", parent=ss["Heading1"], fontSize=14, spaceBefore=14, spaceAfter=7,
                         textColor=colors.HexColor("#1a5670")),
    "h3": ParagraphStyle("H3", parent=ss["Heading2"], fontSize=11.5, spaceBefore=10, spaceAfter=4,
                         textColor=colors.HexColor("#2c3e50")),
    "p": ParagraphStyle("P", parent=ss["BodyText"], fontSize=9.8, leading=14.5, spaceAfter=7,
                        alignment=4),
    "b": ParagraphStyle("B", parent=ss["BodyText"], fontSize=9.8, leading=14, spaceAfter=4,
                        leftIndent=14, bulletIndent=4),
    "quote": ParagraphStyle("Q", parent=ss["BodyText"], fontSize=9.5, leading=14, spaceAfter=8,
                            leftIndent=18, rightIndent=10, textColor=colors.HexColor("#444444"),
                            borderColor=colors.HexColor("#bcd4de"), borderWidth=0,
                            borderPadding=4, fontName="Helvetica-Oblique"),
    "code": ParagraphStyle("C", parent=ss["BodyText"], fontName="Courier", fontSize=8.2,
                           leading=11.5, leftIndent=12, spaceAfter=3,
                           textColor=colors.HexColor("#20303a")),
    "hueco": ParagraphStyle("Hu", parent=ss["BodyText"], fontSize=9, alignment=TA_CENTER,
                            textColor=colors.HexColor("#9a9a9a"), spaceBefore=8, spaceAfter=12,
                            fontName="Helvetica-Oblique"),
}


def esc(s):
    return (str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


flow = []
primero = True
for tipo, val in C:
    if tipo == "h1":
        if not primero:
            flow.append(PageBreak())
        primero = False
        flow.append(Paragraph(esc(val), E["h1"]))
    elif tipo in ("h2", "h3", "p"):
        flow.append(Paragraph(esc(val), E[tipo]))
    elif tipo == "b":
        flow.append(Paragraph(esc(val), E["b"], bulletText="•"))
    elif tipo == "quote":
        flow.append(Paragraph(esc(val), E["quote"]))
    elif tipo == "code":
        for ln in str(val).split("\n"):
            flow.append(Paragraph(esc(ln).replace(" ", "&nbsp;"), E["code"]))
        flow.append(Spacer(1, 5))
    elif tipo == "tabla":
        datos = [[Paragraph(f"<b>{esc(c)}</b>" if i == 0 else esc(c), E["p"]) for c in fila]
                 for i, fila in enumerate(val)]
        t = Table(datos, hAlign="LEFT", colWidths=[16.5 * cm / len(val[0])] * len(val[0]))
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8f1f5")),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#b9ccd6")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        flow.append(t)
        flow.append(Spacer(1, 9))
    elif tipo == "hueco":
        t = Table([[Paragraph(f"[ {esc(val)} ]", E["hueco"])]], colWidths=[16.5 * cm],
                  rowHeights=[2.6 * cm])
        t.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor("#cccccc")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        flow.append(t)
        flow.append(Spacer(1, 10))


def pie(canvas, doc_):
    canvas.saveState()
    canvas.setFont("Helvetica", 7.5)
    canvas.setFillColor(colors.HexColor("#8a8a8a"))
    canvas.drawString(2 * cm, 1.3 * cm, "Dossier de evidencias — Milestone 2 · COT-2026-TUV-001 · corte final 16-ago-2026")
    canvas.drawRightString(19.4 * cm, 1.3 * cm, f"pág. {doc_.page}")
    canvas.restoreState()


ruta_pdf = SALIDA / "DOSSIER-EVIDENCIAS-MILESTONE2.pdf"
SimpleDocTemplate(str(ruta_pdf), pagesize=LETTER, topMargin=2 * cm, bottomMargin=2 * cm,
                  leftMargin=2.2 * cm, rightMargin=2.2 * cm,
                  title="Dossier de evidencias — Milestone 2",
                  author="Equipo Plogy").build(flow, onFirstPage=pie, onLaterPages=pie)
print(f"PDF  -> {ruta_pdf}")
print(f"bloques renderizados: {len(C)}")
